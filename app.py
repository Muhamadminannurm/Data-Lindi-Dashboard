import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
from scipy.stats import pearsonr
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
import joblib
import io
import base64
import plotly.graph_objects as go
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =============================================================================
# 1. KONFIGURASI HALAMAN & CSS
# =============================================================================
st.set_page_config(
    page_title="Leachate Prediction",
    page_icon="💧",
    layout="wide",
    initial_sidebar_state="auto"
)

def get_base64(bin_file):
    with open(bin_file, 'rb') as f:
        data = f.read()
    return base64.b64encode(data).decode()

# --- FUNGSI GENERATOR WORD REPORT (VERSI LENGKAP) ---
def create_word_report(report_type, df_sample, eda_data, ann_data=None):
    doc = Document()
    
    # Header
    title = 'Laporan Analisis EDA' if report_type == 'EDA' else 'Laporan Lengkap Prediksi Lindi'
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Tanggal Generate: {pd.Timestamp.now().strftime("%d-%m-%Y %H:%M")}')
    
    # BAGIAN 1: DATA INPUT (Ada di semua laporan)
    doc.add_heading('1. Sampel Data Input', level=1)
    # Bikin tabel dari dataframe head
    t = doc.add_table(df_sample.shape[0]+1, df_sample.shape[1])
    t.style = 'Table Grid'
    # Header Tabel
    for j, col in enumerate(df_sample.columns):
        t.cell(0, j).text = str(col)
    # Isi Tabel
    for i, row in enumerate(df_sample.itertuples(index=False)):
        for j, val in enumerate(row):
            t.cell(i+1, j).text = str(round(val, 2) if isinstance(val, float) else val)
            
    # BAGIAN 2: EDA (Ada di semua laporan)
    doc.add_heading('2. Exploratory Data Analysis (EDA)', level=1)
    doc.add_paragraph(f"Analisis Fokus Kolom: {eda_data['col']}")
    doc.add_paragraph(f"Insight: {eda_data['insight']}")
    
    for title, fig in eda_data['figs']:
        doc.add_heading(title, level=2)
        buf = io.BytesIO()
        fig.savefig(buf, format='png', bbox_inches='tight')
        buf.seek(0)
        doc.add_picture(buf, width=Inches(5.5))

    # BAGIAN 3: HASIL MODEL (Hanya jika tipe FULL)
    if report_type == 'FULL' and ann_data:
        doc.add_heading('3. Hasil Prediksi AI (ANN)', level=1)
        
        # Tabel Metrik
        doc.add_heading('Metrik Evaluasi', level=2)
        tm = doc.add_table(1, 2)
        tm.style = 'Table Grid'
        tm.cell(0,0).text = "METRIK"; tm.cell(0,1).text = "NILAI"
        for k, v in ann_data['metrics'].items():
            row = tm.add_row().cells
            row[0].text = k; row[1].text = str(v)
            
        doc.add_paragraph(f"\nInsight Model: {ann_data['insight']}")
        
        # Grafik ANN
        for title, fig in ann_data['figs']:
            doc.add_heading(title, level=2)
            buf = io.BytesIO()
            fig.savefig(buf, format='png', bbox_inches='tight')
            buf.seek(0)
            doc.add_picture(buf, width=Inches(6))

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out

def set_style(png_file):
    try:
        bin_str = get_base64(png_file)
        st.markdown(
            f"""
            <style>
            /* --- ANIMASI --- */
            @keyframes slideDownSmooth {{
                0% {{ opacity: 0; transform: translateY(-40px); filter: blur(5px); }}
                100% {{ opacity: 1; transform: translateY(0); filter: blur(0px); }}
            }}
            @keyframes popIn {{
                0% {{ opacity: 0; transform: scale(0.95); }}
                60% {{ opacity: 1; transform: scale(1.02); }}
                100% {{ transform: scale(1); }}
            }}

            /* --- TEKS GLOBAL --- */
            h1, h2, h3, h4, h5, p, label, span, div.stMarkdown, div.stMetricLabel, li, div[data-testid="stDialog"] {{
                color: #e0e0e0 !important;
                font-family: 'Segoe UI', Roboto, Helvetica, Arial, sans-serif;
            }}

            /* --- 1. KEYFRAMES ANIMASI NEON BERJALAN --- */
            @keyframes rgbFlow {{
                0% {{ background-position: 0% 50%; }}
                50% {{ background-position: 100% 50%; }}
                100% {{ background-position: 0% 50%; }}
            }}

            @keyframes slideDownSmooth {{
                0% {{ opacity: 0; transform: translateY(-40px); filter: blur(5px); }}
                100% {{ opacity: 1; transform: translateY(0); filter: blur(0px); }}
            }}
            
            /* --- BACKGROUND UTAMA --- */
            .stApp {{
                background-image: url("data:image/png;base64,{bin_str}");
                background-size: cover;
                background-position: center;
                background-attachment: fixed;
            }}

           /* --- KONTAINER UTAMA (RGB BORDER ONLY) --- */
            .block-container {{
                aposition: relative; /* Wajib ada agar border menempel */
                animation: slideDownSmooth 1.2s cubic-bezier(0.23, 1, 0.32, 1) both;
                
                /* Warna Kaca Gelap Transparan (Tengahnya) */
                background-color: rgba(28, 31, 38, 0.2); 
                
                /* --- PERUBAHAN DISINI (MEMBUAT LEBAR PENUH) --- */
                max-width: 100vw !important;  /* Paksa lebar maksimum seukuran layar */
                width: 100% !important;       /* Paksa lebar 100% */
                padding: 2rem 2rem !important; /* Sesuaikan padding agar tidak terlalu mepet */
                /* ----------------------------------------------- */
                border-radius: 20px;
                padding: 2.5rem 3rem;
                margin-top: 2rem;
                box-shadow: 0 20px 50px rgba(0,0,0,0.8);
                backdrop-filter: blur(10px);
                border: none; /* Border asli dimatikan */
            }}

            .block-container::before {{
                content: "";
                position: absolute;
                inset: 0;
                border-radius: 20px; 
                padding: 3px; /* KETEBALAN GARIS NEON */
                
                /* Warna Pelangi */
                background: linear-gradient(45deg, #00C6FF, #0072FF, #00E676, #00C6FF); 
                background-size: 400% 400%; /* Diperbesar agar animasi jalan terlihat */
                animation: rgbFlow 2s linear infinite; /* Kecepatan jalan */
                
                /* TEKNIK MASKING: Membuat tengahnya bolong sehingga transparansi kaca di layer bawah terlihat */
                -webkit-mask: 
                linear-gradient(#fff 0 0) content-box, 
                linear-gradient(#fff 0 0);
                -webkit-mask-composite: xor;
                mask-composite: exclude;
                
                pointer-events: none; /* Agar tetap bisa klik isi konten */
                z-index: 2; /* Di atas background kaca */
            }}

            /* --- SIDEBAR --- */
            section[data-testid="stSidebar"] {{
                background-color: rgba(20, 23, 28, 0.4); /* Kaca Gelap */
                backdrop-filter: blur(10px);
                border: none; /* Hapus border biasa */
                position: relative; /* Wajib agar neon menempel */
            }}
            section[data-testid="stSidebar"]::before {{
                content: "";
                position: absolute;
                top: 0; 
                right: 0; 
                bottom: 0;
                width: 3px; /* KETEBALAN GARIS NEON */
                
                /* Warna Pelangi Vertikal */
                background-size: 400% 400%;
                animation: rgbFlow 1s linear infinite; /* Animasi Jalan */
                
                z-index: 2;
            }}
            section[data-testid="stSidebar"] .stMarkdown, section[data-testid="stSidebar"] h1 {{ color: #ffffff !important; }}

            /* --- TOMBOL --- */
            div.stButton > button {{
                background: linear-gradient(135deg, #00C6FF 0%, #0072FF 100%);
                color: white !important;
                border: none;
                border-radius: 12px;
                padding: 0.6rem 1.5rem;
                font-weight: 700;
                transition: all 0.4s cubic-bezier(0.23, 1, 0.32, 1);
                box-shadow: 0 4px 15px rgba(0, 114, 255, 0.3);
            }}
            div.stButton > button:hover {{
                transform: translateY(-2px) scale(1.02);
                box-shadow: 0 8px 25px rgba(0, 114, 255, 0.6);
            }}
            
            /* --- WIDGET INPUT --- */
            div[data-baseweb="input"], div[data-baseweb="select"] > div {{
                background-color: rgba(255, 255, 255, 0.05) !important;
                border: 1px solid rgba(255, 255, 255, 0.1) !important;
                color: white !important;
            }}

            /* --- MODAL DIALOG (FIXED DARK GLASS) --- */
            div[role="dialog"] {{
                position: relative; /* Penting agar border neon menempel */
                
                /* Warna Kaca Gelap Transparan */
                background-color: rgba(28, 31, 38, 0.2) !important; 
                backdrop-filter: blur(10px);
                
                border-radius: 20px;
                box-shadow: 0 20px 60px rgba(0,0,0,0.8);
                color: #e0e0e0 !important;
                
                border: none !important;
            }}
            
            div[role="dialog"]::before {{
                content: "";
                position: absolute;
                inset: 0; /* Menempel di ujung */
                border-radius: 20px; 
                padding: 3px; /* KETEBALAN NEON */
                
                /* Warna Pelangi */
                background: linear-gradient(45deg, #00C6FF, #0072FF, #00E676, #00C6FF); 
                background-size: 400% 400%;
                animation: rgbFlow 1s linear infinite;
                
                /* MASKING: Membuat tengah bolong */
                -webkit-mask: 
                   linear-gradient(#fff 0 0) content-box, 
                   linear-gradient(#fff 0 0);
                -webkit-mask-composite: xor;
                mask-composite: exclude;
                
                pointer-events: none;
                z-index: 2;
            }}

            /* --- MENGHILANGKAN TOMBOL 'X' (CLOSE) DI DIALOG --- */
            div[role="dialog"] button[aria-label="Close"] {{
                display: none !important;
            }}
            
            </style>
            """,
            unsafe_allow_html=True
        )
    except FileNotFoundError:
        st.warning("⚠️ File 'TPA.jpg' tidak ditemukan.")

set_style('TPA.jpg')

# =============================================================================
# 2. LOGIKA BACKEND & MODEL
# =============================================================================
try:
    model = joblib.load("model_leachate_ann.pkl")
    scaler_X = joblib.load("scaler_X.pkl")
    MODEL_LOADED = True
    
    try:
        SCALER_FEATURES = list(scaler_X.feature_names_in_)
    except AttributeError:
        SCALER_FEATURES = [
            'TN', 'TX', 'TAVG', 'RH_AVG', 'RR', 'SS', 'FF_X', 'DDD_X', 'FF_AVG',
            'DDD_CAR_C', 'DDD_CAR_E', 'DDD_CAR_NW', 'DDD_CAR_S', 'DDD_CAR_SE', 'DDD_CAR_SW', 'DDD_CAR_W'
        ]
except Exception as e:
    st.error(f"Error System: {str(e)}")
    MODEL_LOADED = False

def scale_features(df_input, scaler_X):
    feature_names = SCALER_FEATURES
    X = df_input[feature_names]
    y = df_input["Lindi"] 
    X_scaled = scaler_X.transform(X)
    return X, y, X_scaled

# =============================================================================
# 3. SIDEBAR NAVIGATION (UPDATED LOGIC)
# =============================================================================
with st.sidebar:
    try:
        st.image("logo.png", use_container_width=True)
    except:
        st.markdown("<h2 style='text-align:center; color:#00C6FF;'>AI LEACHATE</h2>", unsafe_allow_html=True)
    
    st.markdown("---")
    
    # CSS Radio Button
    st.markdown("""
        <style>
        /* Container Dasar Tombol */
        div[role="radiogroup"] label {
            background-color: rgba(255,255,255,0.05) !important;
            color: #ccc !important;
            padding: 12px 15px;
            border-radius: 10px;
            margin-bottom: 8px;
            border: 1px solid rgba(255,255,255,0.05);
            
            /* Transisi Halus untuk efek geser */
            transition: all 0.3s cubic-bezier(0.25, 0.8, 0.25, 1);
            position: relative;
            overflow: hidden;
        }
        
        /* Efek Hover: GESER KANAN 6px */
        div[role="radiogroup"] label:hover {
            background-color: rgba(0, 198, 255, 0.15) !important;
            border-color: #00C6FF;
            
            /* INI BAGIAN ANIMASI GESERNYA */
            transform: translateX(6px); 
            
            color: white !important;
            box-shadow: -3px 0 10px rgba(0, 198, 255, 0.2);
        }
        
        /* Efek Aktif (Terpilih) */
        div[role="radiogroup"] label[data-checked="true"] {
             background: linear-gradient(90deg, rgba(0,198,255,0.2), transparent) !important;
             border-left: 4px solid #00C6FF !important;
             color: #00C6FF !important;
             font-weight: bold;
             
             /* Geser sedikit untuk menandakan aktif */
             transform: translateX(2px);
        }
        </style>
    """, unsafe_allow_html=True)

# --- STATE MANAGEMENT ---
    # Perbarui opsi menu (Tambahkan "🔍 Spesifikasi Model")
    menu_options = ["🏠 Home", "🔍 Spesifikasi Model", "📊 Testing Batch", "📝 Testing Single"]
    
    if "active_menu" not in st.session_state:
        st.session_state["active_menu"] = menu_options[0]

    # --- DIALOG POPUP ---
    @st.dialog("⚠️ Konfirmasi Perpindahan")
    def confirm_switch(new_selection):
        st.markdown(
            f"""
            <div style='text-align: center; color: #e0e0e0;'>
                <p style='font-size: 1.1em;'>
                    Anda akan berpindah ke <br>
                    <b style='color: #00C6FF;'>{new_selection}</b>
                </p>
                <div style='background: rgba(255,255,255,0.05); padding: 15px; border-radius: 10px; font-size: 0.9em; margin: 20px 0; border: 1px solid rgba(255,255,255,0.1);'>
                    ❗ Tindakan ini akan <b>MENGHAPUS (RESET)</b> semua hasil analisis dan data yang sedang tampil saat ini.
                </div>
                <p>Apakah Anda yakin ingin melanjutkan?</p>
            </div>
            """, unsafe_allow_html=True
        )
        
        col_y, col_n = st.columns(2)
        with col_y:
            if st.button("✅ YA, Reset", type="primary", use_container_width=True):
                # --- UPDATE DAFTAR VARIABEL YANG HARUS DIHAPUS ---
                # Masukkan semua key session yang kita pakai di fitur-fitur baru
                keys_to_clear = [
                    'df_processed',       # Dataframe lama
                    'X_train_scaled',     # Data training lama
                    'single_pred',        # Hasil Testing Single (Gauge Chart)
                    'df_processed_result',# Hasil Pre-processing Batch
                    'data_status_index',  # Posisi Radio Button Wizard Batch
                    'active_menu'         # (Opsional, tapi ini biasanya ditimpa)
                ]
                
                for k in keys_to_clear: 
                    if k in st.session_state: 
                        del st.session_state[k]
                
                # Update Halaman Aktif
                st.session_state["active_menu"] = new_selection
                st.rerun()
                
        with col_n:
            if st.button("❌ Batal", use_container_width=True):
                # Kembalikan Radio ke posisi semula
                st.session_state["nav_radio"] = st.session_state["active_menu"]
                st.rerun()

# --- CALLBACK (MODIFIED LOGIC) ---
    def on_nav_change():
        target = st.session_state["nav_radio"]
        current = st.session_state["active_menu"]
        
        # DEFINISI MENU AMAN (Tidak perlu konfirmasi reset)
        # Menu 0: Home
        # Menu 1: Spesifikasi Model (Dokumentasi) -> Tidak ada input user yg perlu direset
        safe_menus = [menu_options[0], menu_options[1]]
        
        # LOGIKA PERPINDAHAN
        if current in safe_menus:
            # Jika dari Home atau Dokumentasi -> Langsung pindah
            st.session_state["active_menu"] = target
            
        elif target != current:
            # Jika dari Testing Batch/Single -> Munculkan Popup Konfirmasi
            confirm_switch(target)

    # --- RENDER RADIO BUTTON ---
    selection = st.radio(
        "NAVIGASI UTAMA",
        menu_options,
        key="nav_radio",
        index=menu_options.index(st.session_state["active_menu"]),
        on_change=on_nav_change
    )
    
    menu = st.session_state["active_menu"]

    st.markdown("---")
    st.markdown(
        """
        <div style='text-align: center; font-size: 12px; color: #888; background: rgba(0,0,0,0.3); padding: 15px; border-radius: 10px;'>
            Developed by:<br>
            <strong style='color: #00C6FF; font-size: 13px; letter-spacing: 0.5px;'>Universitas Muhammadiyah Malang</strong>
        </div>
        """, unsafe_allow_html=True
    )

# =============================================================================
# 4. KONTEN UTAMA
# =============================================================================

# HEADER
st.markdown(
    """
    <div style='text-align: center; margin-bottom: 40px;'>
        <h1 style='margin: 0; font-weight: 800; letter-spacing: 1.5px; text-shadow: 0 4px 8px rgba(0,0,0,0.5);'>PREDIKSI VOLUME AIR LINDI</h1>
        <p style='font-size: 16px; opacity: 0.7; color: #00C6FF !important; font-weight: 500; letter-spacing: 2px; text-transform: uppercase;'>Artificial Neural Network Implementation</p>
        <div style='height: 3px; width: 60px; background: #00C6FF; margin: 20px auto; border-radius: 2px; box-shadow: 0 0 10px #00C6FF;'></div>
    </div>
    """, unsafe_allow_html=True
)

# -----------------------------------------------------------------------------
# HALAMAN: HOME
# -----------------------------------------------------------------------------
if menu == "🏠 Home":
    # Konten Home (Placeholder)
    st.markdown(
        """
        <div style='text-align: center; padding: 20px;'>
            <h2 style='color: #e0e0e0;'>Selamat Datang di Sistem Cerdas TPA</h2>
            <p style='font-size: 18px; color: #ccc; margin-top: 10px;'>
                Sistem ini dirancang untuk membantu memprediksi volume air lindi (leachate) 
                di Tempat Pembuangan Akhir menggunakan algoritma Jaringan Syaraf Tiruan (ANN).
            </p>
            <br>
            <div style='background: rgba(255,255,255,0.05); padding: 20px; border-radius: 15px; border: 1px solid rgba(255,255,255,0.1); display: inline-block; text-align: left;'>
                <h4 style='color: #00C6FF; margin-bottom: 15px;'>🔍 Fitur Utama:</h4>
                <ul style='list-style-type: none; padding: 0; color: #ddd;'>
                    <li style='margin-bottom: 10px;'>📊 <b>Testing Batch:</b> Evaluasi model dengan dataset besar (CSV).</li>
                    <li style='margin-bottom: 10px;'>📝 <b>Testing Single:</b> Prediksi harian cepat dengan input manual.</li>
                    <li style='margin-bottom: 10px;'>📈 <b>Visualisasi:</b> Grafik interaktif Time Series dan Scatter Plot.</li>
                </ul>
            </div>
            <br><br>
            <p style='font-size: 14px; color: #888;'>Silakan pilih menu di sidebar kiri untuk memulai analisis.</p>
        </div>
        """, 
        unsafe_allow_html=True
    )

# -----------------------------------------------------------------------------
# SKENARIO 2: TESTING BATCH (FINAL LOGIC - AUTO DETECT & PRE-PROCESSING TOOL)
# -----------------------------------------------------------------------------
elif menu == "📊 Testing Batch":
    st.markdown("### 📂 Analisis Data Batch")
    
    # --- HELPER: Bikin Grafik Khusus Laporan (Clean Style) ---
    def create_clean_fig(x, y, title, type='line', y2=None, xlabel="", ylabel=""):
        fig, ax = plt.subplots(figsize=(8, 4))
        ax.set_title(title, color='black', fontsize=12, fontweight='bold')
        ax.set_xlabel(xlabel, color='black'); ax.set_ylabel(ylabel, color='black')
        ax.tick_params(colors='black')
        for spine in ax.spines.values(): spine.set_edgecolor('black')
        
        if type == 'hist': ax.hist(x, bins=20, color='#1f77b4', edgecolor='black', alpha=0.7)
        elif type == 'scatter': ax.scatter(x, y, color='#2ca02c', alpha=0.6)
        elif type == 'line_compare': 
            ax.plot(x.index, x.values, label='Aktual', color='green')
            ax.plot(x.index, y2, label='Prediksi', color='blue', linestyle='--'); ax.legend()
        elif type == 'scatter_valid':
            ax.scatter(x, y, color='orange', alpha=0.6)
            ax.plot([x.min(), x.max()], [x.min(), x.max()], 'k--', lw=2)
        elif type == 'bar_error': ax.bar(range(len(x)), x, color='red', alpha=0.8)
        
        plt.tight_layout(); return fig

    # --- STATE MANAGEMENT LOGIC ---
    # Logika Tabulasi Status Data
    # Kita gunakan Radio Button, tapi logikanya bisa dipaksa pindah jika salah upload
    
    # Default state untuk radio button
    if 'data_status_index' not in st.session_state:
        st.session_state['data_status_index'] = 0 # Default: Siap Upload

    # Fungsi callback untuk memindahkan radio button (PENTING)
    def set_status_not_ready():
        st.session_state['data_status_index'] = 1 # Pindah ke "Belum Siap"

    # RADIO BUTTON STATUS
    data_status_opts = ["✅ Data Siap (Sudah Pre-processed)", "❌ Data Belum Siap (Data Mentah)"]
    status_selection = st.radio(
        "Status Data Anda:", 
        data_status_opts, 
        index=st.session_state['data_status_index'],
        key="status_radio_widget" # Key unik agar tidak conflict
    )

    # -------------------------------------------------------------------------
    # SKENARIO A: DATA BELUM SIAP (TOOL PRE-PROCESSING)
    # -------------------------------------------------------------------------
    if status_selection == "❌ Data Belum Siap (Data Mentah)":
        st.info("🛠️ **Tool Pre-processing:** Upload data mentah Anda (harus ada kolom `DDD_CAR`), sistem akan melakukan *One-Hot Encoding* otomatis.")
        
        col_up, col_help = st.columns([2, 1])
        with col_up:
            uploaded_raw = st.file_uploader("Upload File Mentah (Excel/CSV)", type=["xlsx", "csv"])
        with col_help:
            st.markdown("""
            **Panduan Kolom:**
            Pastikan ada kolom `DDD_CAR` (Arah Angin Kategori) untuk diproses menjadi angka (One-Hot Encoding).
            """)

        if uploaded_raw:
            try:
                # Load Data Mentah
                if uploaded_raw.name.endswith('.csv'):
                    df_raw = pd.read_csv(uploaded_raw)
                else:
                    df_raw = pd.read_excel(uploaded_raw)
                
                with st.expander("👁️ Pratinjau Data Mentah"):
                    st.dataframe(df_raw.head(), use_container_width=True)

                # Tombol Proses
                if st.button("⚡ Proses Data (Encoding Otomatis)", type="primary"):
                    if 'DDD_CAR' in df_raw.columns:
                        with st.spinner("Sedang memproses One-Hot Encoding..."):
                            # Logic One-Hot Encoding
                            df_onehot = pd.get_dummies(df_raw['DDD_CAR'], prefix='DDD_CAR', dtype=int)
                            df_processed = pd.concat([df_raw, df_onehot], axis=1)
                            df_processed = df_processed.drop(columns=['DDD_CAR']) # Hapus kolom asal
                            
                            # Simpan ke session untuk download
                            st.session_state['df_processed_result'] = df_processed
                            
                            st.success("✅ Pre-Proses Berhasil! Silakan unduh data di bawah ini.")
                    else:
                        st.error("❌ Kolom 'DDD_CAR' tidak ditemukan. Pastikan nama kolom sesuai.")

                # Tampilkan Tombol Download jika sudah diproses
                if 'df_processed_result' in st.session_state:
                    df_final = st.session_state['df_processed_result']
                    st.markdown("---")
                    st.markdown("#### ⬇️ Unduh Data Hasil Proses")
                    
                    c_down1, c_down2 = st.columns(2)
                    
                    # Download CSV
                    csv_data = df_final.to_csv(index=False).encode('utf-8')
                    c_down1.download_button("📄 Unduh Format CSV", csv_data, "DATA_PROSES_LINDI.csv", "text/csv", use_container_width=True)
                    
                    # Download Excel
                    # Butuh buffer BytesIO untuk Excel
                    buffer = io.BytesIO()
                    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                        df_final.to_excel(writer, index=False, sheet_name='Data_Proses')
                    
                    c_down2.download_button("📗 Unduh Format Excel", buffer.getvalue(), "DATA_PROSES_LINDI.xlsx", 
                                            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
            
            except Exception as e:
                st.error(f"Gagal membaca file: {e}")

    # -------------------------------------------------------------------------
    # SKENARIO B: DATA SIAP (ANALISIS AI)
    # -------------------------------------------------------------------------
    else: # Jika Status == "Data Siap"
        f = st.file_uploader("Upload Data Bersih (Siap Olah)", type=["csv", "xlsx"])
        
        if f and MODEL_LOADED:
            # 1. BACA FILE
            df = pd.read_csv(f) if f.name.endswith('.csv') else pd.read_excel(f)
            
            # 2. VALIDASI KOLOM MODEL (AUTO-DETECT LOGIC)
            # Cek apakah kolom-kolom wajib (termasuk hasil encoding DDD_CAR_...) ada
            required_cols = ['TN', 'TX', 'TAVG', 'RH_AVG', 'RR', 'SS', 'FF_X', 'DDD_X', 'FF_AVG']
            # Cek minimal satu kolom encoding arah angin ada (tanda sudah di encoding)
            encoding_check = any(col.startswith('DDD_CAR_') for col in df.columns)
            basic_check = all(col in df.columns for col in required_cols)
            
            if not basic_check or not encoding_check:
                st.error("⚠️ **Format Data Tidak Sesuai!**")
                st.warning("Sepertinya data Anda belum lengkap atau belum melalui tahap *Encoding* (tidak ada kolom `DDD_CAR_...` atau parameter cuaca hilang).")
                st.markdown("Sistem akan mengalihkan Anda ke menu **Data Belum Siap** untuk melakukan pre-processing.")
                
                # Tombol Redirect Manual (Streamlit tidak bisa auto-rerun paksa radio button dengan mudah tanpa trigger)
                if st.button("🔄 Pindah ke Menu Pre-processing"):
                     st.session_state['data_status_index'] = 1 # Set index ke "Belum Siap"
                     st.rerun()
                st.stop() # Berhenti disini
                
            if 'Lindi' not in df.columns: st.error("❌ Wajib ada kolom 'Lindi'"); st.stop()
            
            # --- JIKA LOLOS VALIDASI: LANJUT ANALISIS ---
            
            # 3. FILTER
            st.markdown("---")
            s, e = st.slider("Rentang Data", 0, len(df)-1, (0, len(df)-1))
            df = df.iloc[s:e+1]
            
            with st.expander("📄 Preview Data"): st.dataframe(df, use_container_width=True)

            # 4. TABS
            tab_eda, tab_ai = st.tabs(["📊 Tab 1: EDA", "🚀 Tab 2: Prediksi AI"])
            
            # --- TAB 1: EDA ---
            with tab_eda:
                cols = df.select_dtypes(include=np.number).columns.tolist()
                c_sel1, c_sel2 = st.columns(2)
                with c_sel1: sel_col = st.selectbox("Fitur (X):", cols, index=cols.index('RR') if 'RR' in cols else 0)
                with c_sel2: target_col = st.selectbox("Target (Y):", cols, index=cols.index('Lindi') if 'Lindi' in cols else 0)
                
                c1, c2 = st.columns(2)
                with c1:
                    fig_h, ax_h = plt.subplots(figsize=(6,4)); fig_h.patch.set_alpha(0); ax_h.patch.set_alpha(0)
                    ax_h.hist(df[sel_col], bins=20, color='#00C6FF', edgecolor='white'); ax_h.axis('off'); ax_h.set_title(f"Hist {sel_col}", color='white')
                    st.pyplot(fig_h)
                with c2:
                    fig_s, ax_s = plt.subplots(figsize=(6,4)); fig_s.patch.set_alpha(0); ax_s.patch.set_alpha(0)
                    ax_s.scatter(df[sel_col], df[target_col], c='#00E676', alpha=0.6); ax_s.set_ylabel(target_col, color='white'); ax_s.tick_params(colors='white'); 
                    for sp in ax_s.spines.values(): sp.set_edgecolor('#555')
                    st.pyplot(fig_s)
                
                corr = df[sel_col].corr(df[target_col])
                insight = f"Korelasi {sel_col} vs {target_col}: {corr:.2f}"
                st.info(f"💡 {insight}")
                
                # Download Report EDA
                try:
                    r1 = create_clean_fig(df[sel_col], None, f"Dist {sel_col}", 'hist', xlabel=sel_col)
                    r2 = create_clean_fig(df[sel_col], df[target_col], f"Korelasi", 'scatter', xlabel=sel_col, ylabel=target_col)
                    doc = create_word_report("EDA", df.head(), {'col':sel_col, 'insight':insight, 'figs':[("Hist",r1),("Scatter",r2)]})
                    st.download_button("📘 Download Report EDA", doc, "Laporan_EDA.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                except: pass

            # --- TAB 2: PREDIKSI ---
            with tab_ai:
                st.markdown("### 🚀 Hasil Evaluasi Statistik")
                X, y, Xs = scale_features(df, scaler_X)
                pred = model.predict(Xs)
                
                mse = mean_squared_error(y, pred); rmse = np.sqrt(mse); mae = mean_absolute_error(y, pred)
                r2 = r2_score(y, pred); pr, p_val = pearsonr(y.values.flatten(), pred.flatten())
                
                m1, m2, m3 = st.columns(3); m1.metric("MSE", f"{mse:.2f}"); m2.metric("RMSE", f"{rmse:.2f}"); m3.metric("MAE", f"{mae:.2f}")
                m4, m5, m6 = st.columns(3); m4.metric("R2 Score", f"{r2:.4f}"); m5.metric("Pearson r", f"{pr:.4f}"); m6.metric("P-Value", f"{p_val:.2e}")
                
                st.markdown("""
                <div style='background: rgba(255, 255, 255, 0.05); padding: 15px; border-radius: 10px; font-size: 0.9em; margin: 20px 0; border-left: 4px solid #00C6FF;'>
                    <strong>💡 Interpretasi Hasil:</strong>
                    <ul style='margin-top: 5px; margin-bottom: 0;'>
                        <li><b>MSE/RMSE/MAE:</b> Mengukur rata-rata kesalahan prediksi. Semakin <b>kecil</b> nilainya, semakin akurat model.</li>
                        <li><b>R² Score:</b> Seberapa baik model menjelaskan variasi data (Mendekati 1.0 = Sempurna).</li>
                        <li><b>Pearson's r:</b> Tingkat korelasi linear antara aktual dan prediksi (Mendekati 1.0 = Sangat Kuat).</li>
                        <li><b>P-value:</b> Validitas statistik. Nilai < 0.05 menunjukkan korelasi signifikan (bukan kebetulan).</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)
                st.markdown("#### 1. Time Series")
                fig1, ax1 = plt.subplots(figsize=(10,4)); fig1.patch.set_alpha(0); ax1.patch.set_alpha(0)
                ax1.plot(y.values, color='#00E676', label='Aktual'); ax1.plot(pred, color='#00C6FF', linestyle='--', label='Prediksi')
                ax1.legend(facecolor='#1c1f26', labelcolor='white'); ax1.tick_params(colors='white'); 
                for s in ax1.spines.values(): s.set_visible(False)
                st.pyplot(fig1)
                
                st.markdown("#### 2. Scatter Validasi")
                fig2, ax2 = plt.subplots(figsize=(8,6)); fig2.patch.set_alpha(0); ax2.patch.set_alpha(0)
                ax2.scatter(y, pred, color='#FFC107', alpha=0.6); ax2.plot([y.min(), y.max()], [y.min(), y.max()], 'w--')
                ax2.set_ylabel("Prediksi", color='white'); ax2.tick_params(colors='white'); 
                for s in ax2.spines.values(): s.set_edgecolor('#555')
                st.pyplot(fig2)
                
                st.markdown("#### 3. Error Absolut")
                fig3, ax3 = plt.subplots(figsize=(10,4)); fig3.patch.set_alpha(0); ax3.patch.set_alpha(0)
                errs = np.abs(y.values.flatten() - pred.flatten())
                ax3.bar(range(len(errs)), errs, color='#FF5252', alpha=0.8); ax3.set_ylabel("Error", color='white'); ax3.tick_params(colors='white')
                for s in ax3.spines.values(): s.set_visible(False)
                st.pyplot(fig3)
                
                st.markdown("---")
                try:
                    # Generate Clean Figs for Report
                    rc1 = create_clean_fig(y, None, "Time Series", 'line_compare', y2=pred, xlabel="Index", ylabel="Vol")
                    rc2 = create_clean_fig(y, pred, "Validasi Scatter", 'scatter_valid', xlabel="Aktual", ylabel="Prediksi")
                    rc3 = create_clean_fig(errs, None, "Error Plot", 'bar_error', xlabel="Index", ylabel="Error")
                    
                    ann_pkg = {
                        'metrics': {"MSE":f"{mse:.2f}", "R2":f"{r2:.4f}", "MAE":f"{mae:.2f}", "P-Val":f"{p_val:.2e}"},
                        'insight': f"Model R2: {r2:.4f}. Error rata-rata: {mae:.2f}",
                        'figs': [("Time Series", rc1), ("Scatter", rc2), ("Error", rc3)]
                    }
                    # EDA pkg dummy for full report structure (reuse current tab 1 state logic or empty)
                    # For safety, regenerate basic eda figs
                    re1 = create_clean_fig(df[sel_col], None, f"Hist {sel_col}", 'hist')
                    re2 = create_clean_fig(df[sel_col], df[target_col], "Korelasi", 'scatter')
                    eda_pkg_full = {'col':sel_col, 'insight':insight, 'figs':[("Hist", re1), ("Scatter", re2)]}
                    
                    doc_full = create_word_report("FULL", df.head(), eda_pkg_full, ann_pkg)
                    st.download_button("📘 Download Laporan Lengkap (Word)", doc_full, "Laporan_Full.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                except Exception as e: st.error(f"Err Word: {e}")
                
# -----------------------------------------------------------------------------
# SKENARIO 3: TESTING SINGLE (REVISI - SARAN 1)
# -----------------------------------------------------------------------------
elif menu == "📝 Testing Single":
    if not MODEL_LOADED: st.stop()
    
    st.markdown("### 📝 Simulasi Prediksi Harian")
    st.info("Masukkan parameter cuaca harian di bawah ini untuk memprediksi volume lindi.")

    # --- 1. FORM INPUT DENGAN VALIDASI (Agar User Tidak Salah Input) ---
    with st.container():
        st.markdown("<div style='background:rgba(255,255,255,0.02); padding:20px; border-radius:15px; border:1px solid rgba(255,255,255,0.05);'>", unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        
        # Validasi: Suhu dibatasi -10 sampai 60 derajat (Logis)
        with c1:
            tn = st.number_input("TN (Suhu Min °C)", min_value=-10.0, max_value=60.0, value=22.2, step=0.1)
            tx = st.number_input("TX (Suhu Maks °C)", min_value=-10.0, max_value=60.0, value=31.0, step=0.1)
            tavg = st.number_input("TAVG (Rata-rata °C)", min_value=-10.0, max_value=60.0, value=25.2, step=0.1)
        
        # Validasi: Kelembaban max 100%, Hujan max 500mm
        with c2:
            rh = st.number_input("RH_AVG (Kelembaban %)", min_value=0.0, max_value=100.0, value=83.3, step=0.1)
            rr = st.number_input("RR (Curah Hujan mm)", min_value=0.0, max_value=500.0, value=10.1, step=0.1)
            ss = st.number_input("SS (Sinar Matahari Jam)", min_value=0.0, max_value=24.0, value=5.2, step=0.1)
        
        # Validasi: Angin & Arah
        with c3:
            ffx = st.number_input("FF_X (Angin Max m/s)", min_value=0.0, max_value=100.0, value=2.7, step=0.1)
            dddx = st.number_input("DDD_X (Arah Max °)", min_value=0.0, max_value=360.0, value=198.0, step=1.0)
            ffavg = st.number_input("FF_AVG (Angin Rata m/s)", min_value=0.0, max_value=100.0, value=1.4, step=0.1)
        
        st.markdown("---")
        wind_opts = ['DDD_CAR_C', 'DDD_CAR_E', 'DDD_CAR_NW', 'DDD_CAR_S', 'DDD_CAR_SE', 'DDD_CAR_SW', 'DDD_CAR_W']
        dom_dir = st.radio("Arah Angin Dominan:", wind_opts, horizontal=True)
        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)

    # --- 2. TOMBOL EKSEKUSI ---
    if st.button("🔍 HITUNG PREDIKSI", use_container_width=True):
        try:
            # Persiapan Data
            feat = {'TN':tn, 'TX':tx, 'TAVG':tavg, 'RH_AVG':rh, 'RR':rr, 'SS':ss, 'FF_X':ffx, 'DDD_X':dddx, 'FF_AVG':ffavg}
            for w in wind_opts: feat[w] = 0
            feat[dom_dir] = 1
            
            # Prediksi Model
            df_in = pd.DataFrame([feat])
            # Pastikan urutan kolom sesuai scaler
            df_in = df_in[SCALER_FEATURES] 
            X_sc = scaler_X.transform(df_in)
            pred_val = model.predict(X_sc)[0]
            
            # Simpan ke session state agar tidak hilang saat klik simulasi
            st.session_state['single_pred'] = pred_val
            
        except Exception as e:
            st.error(f"Terjadi kesalahan input: {str(e)}")

    # --- 3. OUTPUT VISUALISASI (GAUGE CHART) ---
    if 'single_pred' in st.session_state:
        val = st.session_state['single_pred']
        
        st.markdown("---")
        
        # Tampilan Angka Besar
        st.markdown(f"""
        <div style='text-align: center; background: linear-gradient(135deg, rgba(0, 198, 255, 0.1), rgba(0,0,0,0)); padding: 20px; border-radius: 20px; border: 1px solid #00C6FF; margin-bottom: 20px;'>
            <h3 style='margin:0; color:#00C6FF; font-size: 1.2em;'>ESTIMASI VOLUME LINDI</h3>
            <h1 style='font-size: 3.5em; margin: 5px 0; color: white; text-shadow: 0 0 10px rgba(0,198,255,0.5);'>{val:.2f}</h1>
            <p style='color: #ccc; margin:0;'>Meter Kubik (m³) / Hari</p>
        </div>
        """, unsafe_allow_html=True)
        
        # FITUR RAHASIA (SIMULASI KAPASITAS)
        # Menggunakan st.expander agar UI tetap bersih (Hidden by default)
        with st.expander("⚙️ Simulasi Kapasitas Penampungan (Decision Support System)"):
            st.info("Masukkan kapasitas kolam penampungan aktual hari ini untuk melihat status keamanan.")
            
            # Input Kapasitas (User Input)
            col_cap, col_chart = st.columns([1, 2])
            
            with col_cap:
                cap_input = st.number_input("Kapasitas Kolam (m³)", min_value=1.0, max_value=5000.0, value=100.0, step=10.0)
                
                # Logika Persentase
                percent_filled = (val / cap_input) * 100
                
                # Logika Warna (Traffic Light)
                if percent_filled <= 50:
                    status_color = "#00E676" # Hijau (Aman)
                    status_text = "AMAN"
                elif percent_filled <= 80:
                    status_color = "#FFC107" # Kuning (Waspada)
                    status_text = "WASPADA"
                else:
                    status_color = "#FF5252" # Merah (Bahaya)
                    status_text = "BAHAYA"
                    
                st.markdown(f"""
                <div style='margin-top: 20px; padding: 15px; background: {status_color}20; border-left: 5px solid {status_color}; border-radius: 5px;'>
                    <strong style='color: {status_color}; font-size: 1.2em;'>STATUS: {status_text}</strong><br>
                    <span style='color: #ccc; font-size: 0.9em;'>Terisi: {percent_filled:.1f}%</span>
                </div>
                """, unsafe_allow_html=True)

            with col_chart:
                # Membuat Gauge Chart (Speedometer)
                fig_gauge = go.Figure(go.Indicator(
                    mode = "gauge+number",
                    value = val,
                    domain = {'x': [0, 1], 'y': [0, 1]},
                    title = {'text': "Volume vs Kapasitas", 'font': {'size': 14, 'color': "white"}},
                    gauge = {
                        'axis': {'range': [None, cap_input * 1.2], 'tickwidth': 1, 'tickcolor': "white"},
                        'bar': {'color': status_color}, # Warna Bar mengikuti status
                        'bgcolor': "rgba(255,255,255,0.05)",
                        'borderwidth': 1,
                        'bordercolor': "#555",
                        'threshold': {
                            'line': {'color': "red", 'width': 4},
                            'thickness': 0.75,
                            'value': cap_input # Garis merah di batas kapasitas
                        }
                    }
                ))
                
                # Konfigurasi Layout agar transparan dan pas
                fig_gauge.update_layout(
                    paper_bgcolor = "rgba(0,0,0,0)",
                    font = {'color': "white", 'family': "Segoe UI"},
                    margin=dict(l=30, r=30, t=30, b=0),
                    height=250
                )
                st.plotly_chart(fig_gauge, use_container_width=True)

        # -----------------------------------------------------------------------------
# MENU BARU: SPESIFIKASI MODEL (DOKUMENTASI HKI)
# -----------------------------------------------------------------------------
elif menu == "🔍 Spesifikasi Model":
    st.markdown("### 📘 Dokumentasi Teknis & Spesifikasi Model")
    st.info("Halaman ini menyajikan detail arsitektur kecerdasan buatan dan variabel data yang digunakan dalam sistem.")

    # Gunakan Tabs agar rapi
    tab_spec1, tab_spec2, tab_spec3 = st.tabs(["🧠 Arsitektur ANN", "📊 Metrik Performa", "📚 Kamus Data"])
    
    # --- TAB 1: ARSITEKTUR ---
    with tab_spec1:
        st.markdown("#### 1. Metode Algoritma")
        st.write("""
        Sistem ini menggunakan algoritma **Artificial Neural Network (ANN)** dengan metode pembelajaran 
        *Backpropagation*. Model dilatih untuk mengenali pola non-linear antara parameter cuaca 
        dan volume air lindi.
        """)
        
        c_arch1, c_arch2 = st.columns([1, 1])
        
        with c_arch1:
            st.markdown("**Konfigurasi Jaringan (Layer):**")
            st.code("""
Input Layer  : 16 Neuron (Fitur Cuaca & Arah Angin)
Hidden Layer 1 : 64 Neuron (Aktivasi: ReLU)
Hidden Layer 2 : 32 Neuron (Aktivasi: ReLU)
Output Layer : 1 Neuron (Aktivasi: Linear)
Optimizer    : Adam (Adaptive Moment Estimation)
Loss Function: Mean Squared Error (MSE)
            """, language="yaml")
            
        with c_arch2:
            st.markdown("**Visualisasi Arsitektur:**")
            # Placeholder Image (Jika Anda punya gambar diagram ANN sendiri, ganti nama file-nya)
            # Jika tidak punya, kode ini akan menampilkan teks placeholder rapi.
            st.warning("Diagram Arsitektur: Input Layer ➡ Hidden Layers ➡ Output Layer")
            
            # Tips: Anda bisa screenshot diagram model dari jurnal/skripsi Anda, 
            # simpan sebagai 'ann_diagram.png' dan uncomment baris bawah ini:
            # st.image("ann_diagram.png", caption="Arsitektur Model ANN Leachate Pro", use_container_width=True)

    # --- TAB 2: METRIK PERFORMA (STATIC DATA DARI TRAINING) ---
    with tab_spec2:
        st.markdown("#### 2. Evaluasi Model (Tahap Training)")
        st.write("Berikut adalah performa model saat diuji menggunakan data validasi (Historical Data):")
        
        # Contoh data statis (Sesuaikan dengan hasil terbaik skripsi Anda)
        col_m1, col_m2, col_m3 = st.columns(3)
        col_m1.metric("Akurasi (R² Score)", "0.8278", help="Nilai mendekati 1.0 menunjukkan model sangat akurat")
        col_m2.metric("Mean Squared Error", "5134.97", help="Rata-rata kesalahan kuadrat")
        col_m3.metric("Korelasi (Pearson)", "0.91", help="Tingkat hubungan linear prediksi vs aktual")
        
        st.caption("*Data evaluasi berdasarkan dataset training tahun 2023-2024.")

    # --- TAB 3: KAMUS DATA ---
    with tab_spec3:
        st.markdown("#### 3. Variabel Input (Fitur)")
        st.write("Penjelasan kode variabel yang digunakan dalam dataset:")
        
        # Buat Dataframe Kamus Data
        data_dict = pd.DataFrame([
            ["TN", "Temperatur Minimum (°C)", "Suhu udara terendah yang tercatat dalam 24 jam."],
            ["TX", "Temperatur Maksimum (°C)", "Suhu udara tertinggi yang tercatat dalam 24 jam."],
            ["TAVG", "Temperatur Rata-rata (°C)", "Rata-rata suhu harian."],
            ["RH_AVG", "Kelembaban Rata-rata (%)", "Persentase uap air di udara."],
            ["RR", "Curah Hujan (mm)", "Intensitas hujan harian."],
            ["SS", "Lama Penyinaran (Jam)", "Durasi matahari bersinar cerah dalam sehari."],
            ["FF_X", "Kecepatan Angin Maks (m/s)", "Kecepatan angin tertinggi sesaat."],
            ["DDD_X", "Arah Angin Maks (Derajat)", "Arah dari mana angin bertiup saat kecepatan maksimum."],
            ["FF_AVG", "Kecepatan Angin Rata-rata", "Rata-rata kecepatan angin harian."],
            ["Lindi", "Volume Air Lindi (m³)", "Target Prediksi: Volume limbah cair yang dihasilkan TPA."]
        ], columns=["Kode Variabel", "Nama Parameter", "Definisi"])
        
        st.table(data_dict)

